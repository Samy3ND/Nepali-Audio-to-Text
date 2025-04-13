import ttkbootstrap as tb
from ttkbootstrap.constants import *
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext
from pydub import AudioSegment
import speech_recognition as sr
from docx import Document
from reportlab.pdfgen import canvas
import threading

def transcribe_audio():
    file_path = filedialog.askopenfilename(filetypes=[("Audio files", "*.m4a *.wav *.mp3")])
    if not file_path:
        return

    log_label.config(text="🔄 Converting and Transcribing... Please wait.")
    transcribe_btn.config(state="disabled")
    save_word_btn.pack_forget()
    save_pdf_btn.pack_forget()

    def process():
        try:
            audio = AudioSegment.from_file(file_path)
            audio = audio.set_channels(1).set_frame_rate(16000).set_sample_width(2)
            clean_path = "clean.wav"
            audio.export(clean_path, format="wav")

            recognizer = sr.Recognizer()
            with sr.AudioFile(clean_path) as source:
                audio_data = recognizer.record(source)

            text = recognizer.recognize_google(audio_data, language="ne-NP")

            text_box.delete("1.0", tk.END)
            text_box.insert(tk.END, text)
            log_label.config(text="✅ Transcription Complete!")

            # Show the download buttons above the text area
            save_word_btn.pack(side=tk.LEFT, padx=10)
            save_pdf_btn.pack(side=tk.LEFT, padx=10)
        except sr.RequestError:
            log_label.config(text=" API error. Please check your internet.")
        except sr.UnknownValueError:
            log_label.config(text=" Couldn't understand the audio.")
        except Exception as e:
            log_label.config(text=f" Error: {e}")
        finally:
            transcribe_btn.config(state="normal")

    threading.Thread(target=process).start()

def save_as_docx():
    content = text_box.get("1.0", tk.END).strip()
    if not content:
        messagebox.showwarning("No Content", "Please transcribe an audio file first.")
        return
    path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
    if path:
        doc = Document()
        doc.add_heading("Audio Transcription", 0)
        doc.add_paragraph(content)
        doc.save(path)
        messagebox.showinfo("Saved", f"Saved as {path}")

def save_as_pdf():
    content = text_box.get("1.0", tk.END).strip()
    if not content:
        messagebox.showwarning("No Content", "Please transcribe an audio file first.")
        return
    path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF File", "*.pdf")])
    if path:
        pdf = canvas.Canvas(path)
        pdf.setFont("Helvetica", 12)
        y = 800
        for line in content.split('\n'):
            pdf.drawString(50, y, line)
            y -= 20
            if y < 50:
                pdf.showPage()
                y = 800
        pdf.save()
        messagebox.showinfo("Saved", f"Saved as {path}")

app = tb.Window(themename="litera")
app.title(" Nepali Audio to Text Converter")
app.geometry("750x650")

title_label = tb.Label(app, text="Nepali Audio Transcription Tool", font=("Helvetica", 20, "bold"))
title_label.pack(pady=15)

transcribe_btn = tb.Button(
    app, text="📤 Upload Audio & Transcribe",
    bootstyle=PRIMARY, width=30,
    command=transcribe_audio
)
transcribe_btn.pack(pady=10)

log_label = tb.Label(app, text="", font=("Helvetica", 11))
log_label.pack(pady=5)

download_frame = tb.Frame(app)
download_frame.pack(pady=5)

save_word_btn = tb.Button(download_frame, text="💾 Save as Word", bootstyle=SUCCESS, width=20, command=save_as_docx)
save_pdf_btn = tb.Button(download_frame, text="📄 Save as PDF", bootstyle=INFO, width=20, command=save_as_pdf)

text_box = scrolledtext.ScrolledText(
    app, wrap=tk.WORD, height=20, font=("Consolas", 11),
    bg="white", fg="black", insertbackground="black"
)
text_box.pack(padx=15, pady=10, fill=tk.BOTH, expand=True)

app.mainloop()
